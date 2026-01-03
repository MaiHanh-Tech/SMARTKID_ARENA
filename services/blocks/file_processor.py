from pypdf import PdfReader
from docx import Document
import re
import os

def doc_file(uploaded_file):
    """
    Đọc nội dung từ file upload hoặc file path
    Hỗ trợ: PDF, DOCX, TXT, MD
    
    Args:
        uploaded_file: 
            - Streamlit UploadedFile object
            - FakeUploadedFile object (có thuộc tính _path)
            - String path
    
    Returns:
        str: Nội dung file (cleaned)
    """
    if not uploaded_file: 
        return ""
    
    try:
        # ✅ CASE 1: Nếu là string path
        if isinstance(uploaded_file, str):
            file_path = uploaded_file
            ext = file_path.split('.')[-1].lower()
            
            if ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
                return clean_text(text)
            
            elif ext == "docx":
                doc = Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                return clean_text(text)
            
            elif ext in ["txt", "md", "html"]:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return clean_text(f.read())
        
        # ✅ CASE 2: Nếu có thuộc tính _path (FakeUploadedFile)
        elif hasattr(uploaded_file, '_path'):
            file_path = uploaded_file._path
            ext = file_path.split('.')[-1].lower()
            
            if ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
                return clean_text(text)
            
            elif ext == "docx":
                doc = Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                return clean_text(text)
            
            elif ext in ["txt", "md", "html"]:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return clean_text(f.read())
        
        # ✅ CASE 3: Streamlit UploadedFile (có method read())
        else:
            ext = uploaded_file.name.split('.')[-1].lower()
            
            if ext == "pdf":
                reader = PdfReader(uploaded_file)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
                return clean_text(text)
                
            elif ext == "docx":
                doc = Document(uploaded_file)
                text = "\n".join([p.text for p in doc.paragraphs])
                return clean_text(text)
                
            elif ext in ["txt", "md", "html"]:
                content = uploaded_file.read()
                if isinstance(content, bytes):
                    return clean_text(content.decode('utf-8'))
                return clean_text(content)
                    
    except Exception as e:
        print(f"❌ Lỗi đọc file: {type(e).__name__}: {e}")
        return ""
    
    return ""

def clean_text(text):
    """Làm sạch văn bản (loại bỏ ký tự lạ, khoảng trắng thừa)"""
    if not text: 
        return ""
    
    # Loại bỏ ký tự điều khiển
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    # Chuẩn hóa khoảng trắng
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()
